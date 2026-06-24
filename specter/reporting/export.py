"""Export Markdown reports to PDF and Word (.docx).

The heavy rendering libraries are optional and imported lazily, so the base
install stays light. Install with:

    pip install "ai-specter[pdf]"     # PDF  (fpdf2, pure-Python)
    pip install "ai-specter[docx]"    # Word (python-docx)
    pip install "ai-specter[reports]" # both

A tiny Markdown subset is parsed (headings, paragraphs, bullet/numbered lists,
pipe tables, fenced code, and **bold**/_italic_/`code` inline) — enough for the
reports Specter generates, with no native dependencies.
"""
from __future__ import annotations

import re
from pathlib import Path

_INLINE = re.compile(r"\*\*(.+?)\*\*|`(.+?)`|_(.+?)_")
_HEADING = re.compile(r"^(#{1,6})\s+(.*)$")
_ORDERED = re.compile(r"^\d+\.\s+(.*)$")
_BULLET = re.compile(r"^[-*]\s+(.*)$")
_TABLE_SEP = re.compile(r"^\s*\|?[\s:|-]*-[\s:|-]*\|?\s*$")

# Transliterate characters the PDF core fonts (Latin-1) cannot render.
_ASCII_MAP = {
    "—": "-", "–": "-", "•": "*", "→": "->", "·": "-", "✓": "v",
    "’": "'", "“": '"', "”": '"', "…": "...",
}


class ExportUnavailable(RuntimeError):
    """Raised when an optional export backend (fpdf2 / python-docx) is missing."""


def _plain(text: str) -> str:
    """Strip inline Markdown markers, keeping the inner text."""
    return _INLINE.sub(lambda m: m.group(1) or m.group(2) or m.group(3) or "", text)


def _ascii(text: str) -> str:
    for src, dst in _ASCII_MAP.items():
        text = text.replace(src, dst)
    return text.encode("latin-1", "replace").decode("latin-1")


def _split_row(line: str) -> list[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def _parse_markdown(md: str) -> list[tuple]:
    """Parse the Markdown subset into ('kind', ...payload) blocks."""
    lines = md.splitlines()
    blocks: list[tuple] = []
    para: list[str] = []
    i, n = 0, len(lines)

    def flush() -> None:
        if para:
            blocks.append(("p", " ".join(para).strip()))
            para.clear()

    while i < n:
        raw = lines[i]
        line = raw.strip()
        if line.startswith("```"):
            flush()
            i += 1
            code: list[str] = []
            while i < n and not lines[i].strip().startswith("```"):
                code.append(lines[i])
                i += 1
            i += 1  # consume closing fence
            blocks.append(("code", "\n".join(code)))
            continue
        if not line:
            flush()
            i += 1
            continue
        m = _HEADING.match(line)
        if m:
            flush()
            blocks.append(("h", len(m.group(1)), m.group(2).strip()))
            i += 1
            continue
        if line.startswith("|") and i + 1 < n and _TABLE_SEP.match(lines[i + 1]):
            flush()
            header = _split_row(line)
            i += 2
            rows: list[list[str]] = []
            while i < n and lines[i].strip().startswith("|"):
                rows.append(_split_row(lines[i].strip()))
                i += 1
            blocks.append(("table", header, rows))
            continue
        if _ORDERED.match(line):
            flush()
            items = []
            while i < n and _ORDERED.match(lines[i].strip()):
                items.append(_ORDERED.match(lines[i].strip()).group(1))
                i += 1
            blocks.append(("ol", items))
            continue
        if _BULLET.match(line):
            flush()
            items = []
            while i < n and _BULLET.match(lines[i].strip()):
                items.append(_BULLET.match(lines[i].strip()).group(1))
                i += 1
            blocks.append(("ul", items))
            continue
        para.append(line)
        i += 1
    flush()
    return blocks


def _add_runs(paragraph, text: str) -> None:
    """Add inline-formatted runs (**bold**, `code`, _italic_) to a docx paragraph."""
    pos = 0
    for m in _INLINE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        if m.group(1) is not None:
            paragraph.add_run(m.group(1)).bold = True
        elif m.group(2) is not None:
            run = paragraph.add_run(m.group(2))
            run.font.name = "Courier New"
        elif m.group(3) is not None:
            paragraph.add_run(m.group(3)).italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _to_docx(blocks: list[tuple], path: Path, title: str) -> Path:
    try:
        import docx
        from docx.shared import Pt
    except ImportError as e:  # pragma: no cover - exercised only without the extra
        raise ExportUnavailable(
            "install 'ai-specter[docx]' (python-docx) for Word export") from e

    doc = docx.Document()
    for b in blocks:
        kind = b[0]
        if kind == "h":
            doc.add_heading(_plain(b[2]), level=min(b[1], 4))
        elif kind == "p":
            _add_runs(doc.add_paragraph(), b[1])
        elif kind == "ul":
            for it in b[1]:
                _add_runs(doc.add_paragraph(style="List Bullet"), it)
        elif kind == "ol":
            for it in b[1]:
                _add_runs(doc.add_paragraph(style="List Number"), it)
        elif kind == "code":
            run = doc.add_paragraph().add_run(b[1])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        elif kind == "table":
            header, rows = b[1], b[2]
            table = doc.add_table(rows=1, cols=max(1, len(header)))
            table.style = "Table Grid"
            for j, cell in enumerate(header):
                table.rows[0].cells[j].text = _plain(cell)
            for row in rows:
                cells = table.add_row().cells
                for j in range(len(header)):
                    cells[j].text = _plain(row[j]) if j < len(row) else ""
    doc.save(str(path))
    return path


def _pdf_table(pdf, header: list[str], rows: list[list[str]], width: float) -> None:
    cols = max(1, len(header))
    col_w = width / cols
    cap = max(4, int(col_w / 2))
    pdf.set_font("Helvetica", style="B", size=10)
    pdf.set_x(pdf.l_margin)
    for cell in header:
        pdf.cell(col_w, 7, _ascii(_plain(cell))[:cap], border=1)
    pdf.ln()
    pdf.set_font("Helvetica", size=10)
    for row in rows:
        pdf.set_x(pdf.l_margin)
        for j in range(cols):
            val = _plain(row[j]) if j < len(row) else ""
            pdf.cell(col_w, 6, _ascii(val)[:cap], border=1)
        pdf.ln()
    pdf.ln(2)


def _to_pdf(blocks: list[tuple], path: Path, title: str) -> Path:
    try:
        from fpdf import FPDF
    except ImportError as e:  # pragma: no cover - exercised only without the extra
        raise ExportUnavailable(
            "install 'ai-specter[pdf]' (fpdf2) for PDF export") from e

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    width = pdf.w - pdf.l_margin - pdf.r_margin
    heading_sizes = {1: 18, 2: 14, 3: 12}

    def block_text(text: str, size: int = 11, style: str = "", gap: int = 2) -> None:
        pdf.set_font("Helvetica", style=style, size=size)
        pdf.set_x(pdf.l_margin)  # reset cursor: multi_cell leaves x at the right edge
        pdf.multi_cell(0, 6, _ascii(_plain(text)) or " ")
        pdf.ln(gap)

    for b in blocks:
        kind = b[0]
        if kind == "h":
            block_text(b[2], size=heading_sizes.get(b[1], 11), style="B")
        elif kind == "p":
            block_text(b[1])
        elif kind in ("ul", "ol"):
            pdf.set_font("Helvetica", size=11)
            for k, it in enumerate(b[1], 1):
                marker = "*" if kind == "ul" else f"{k}."
                pdf.set_x(pdf.l_margin)
                pdf.multi_cell(0, 6, _ascii(f"{marker} {_plain(it)}"))
            pdf.ln(2)
        elif kind == "code":
            pdf.set_font("Courier", size=9)
            pdf.set_x(pdf.l_margin)
            pdf.multi_cell(0, 5, _ascii(b[1]) or " ")
            pdf.ln(2)
        elif kind == "table":
            _pdf_table(pdf, b[1], b[2], width)
    pdf.output(str(path))
    return path


_EXPORTERS = {"pdf": _to_pdf, "docx": _to_docx}


def export_markdown_file(md_path: Path, fmt: str) -> Path:
    """Render a Markdown report file to a sibling .pdf or .docx and return its path."""
    fmt = fmt.lower()
    if fmt not in _EXPORTERS:
        raise ExportUnavailable(f"unknown report format '{fmt}' (use pdf or docx)")
    md_path = Path(md_path)
    blocks = _parse_markdown(md_path.read_text())
    title = next((b[2] for b in blocks if b[0] == "h"), md_path.stem)
    out = md_path.with_suffix(".docx" if fmt == "docx" else ".pdf")
    return _EXPORTERS[fmt](blocks, out, title)


def available_export_formats() -> dict[str, bool]:
    """Report which export backends are importable (for `doctor`)."""
    import importlib.util as u
    return {"pdf": u.find_spec("fpdf") is not None,
            "docx": u.find_spec("docx") is not None}
